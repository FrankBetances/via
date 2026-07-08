#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera el Manual de Usuario de VIA+ en formato Word (.docx).

Cada página del manual se rasteriza a alta resolución (reutilizando
`build-pdf.js` con EMIT_PAGES) y se incrusta a página completa A4 en el
documento. El resultado es una réplica visual fiel del PDF que abre en
Microsoft Word y Google Docs.

Uso:
    python3 docs/manual/build-docx.py

Requisitos:
    - Node con Playwright (para build-pdf.js).
    - python-docx  (pip install python-docx).

Variables de entorno opcionales (se pasan a build-pdf.js):
    CHROMIUM_PATH, PLAYWRIGHT_MODULE
"""
import os
import glob
import shutil
import tempfile
import subprocess

from docx import Document
from docx.shared import Mm, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, 'VIA+_Manual_de_Usuario.docx')


def render_pages(dest):
    env = dict(os.environ, EMIT_PAGES=dest)
    subprocess.run(['node', os.path.join(HERE, 'build-pdf.js')], check=True, env=env)
    return sorted(glob.glob(os.path.join(dest, 'p*.png')))


def build_docx(images):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    for attr in ('left_margin', 'right_margin', 'top_margin', 'bottom_margin',
                 'gutter', 'header_distance', 'footer_distance'):
        setattr(sec, attr, Mm(0))

    normal = doc.styles['Normal']
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    normal.font.size = Pt(1)

    for i, img in enumerate(images):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        if i > 0:
            pf.page_break_before = True   # una imagen por página, sin desbordes
        # 296.5 mm de alto: un pelo por debajo de A4 para no forzar una página en blanco
        p.add_run().add_picture(img, width=Mm(210), height=Mm(296.5))

    doc.save(OUT)
    return OUT


def main():
    tmp = tempfile.mkdtemp(prefix='via-manual-pages-')
    try:
        images = render_pages(tmp)
        if not images:
            raise SystemExit('No se generaron imágenes de página.')
        build_docx(images)
        print('DOCX:', OUT, round(os.path.getsize(OUT) / 1024), 'KB')
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


if __name__ == '__main__':
    main()
